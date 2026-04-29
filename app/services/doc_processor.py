import io
from dataclasses import asdict, dataclass, field
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor

from app.core.config import settings


@dataclass
class RunStyleSnapshot:
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    font_name: str | None = None
    font_size_pt: float | None = None
    color_rgb: str | None = None
    all_caps: bool | None = None


@dataclass
class ParagraphSnapshot:
    block_id: str
    text: str
    style_name: str | None
    alignment: str
    left_indent_pt: float | None
    right_indent_pt: float | None
    first_line_indent_pt: float | None
    space_before_pt: float | None
    space_after_pt: float | None
    keep_together: bool | None
    keep_with_next: bool | None
    page_break_before: bool | None
    run_styles: list[RunStyleSnapshot] = field(default_factory=list)


@dataclass
class TableSnapshot:
    block_id: str
    rows: list[list[str]]
    style_name: str | None
    row_count: int
    col_count: int


@dataclass
class ImageSnapshot:
    block_id: str
    image_bytes: bytes
    width_pt: float | None
    height_pt: float | None
    alignment: str


@dataclass
class SectionSnapshot:
    top_margin_pt: float | None
    bottom_margin_pt: float | None
    left_margin_pt: float | None
    right_margin_pt: float | None
    page_width_pt: float | None
    page_height_pt: float | None


@dataclass
class DocumentPackage:
    blocks: list[dict[str, Any]]
    paragraphs: list[ParagraphSnapshot]
    tables: list[TableSnapshot]
    images: list[ImageSnapshot]
    section: SectionSnapshot | None
    style_catalog: list[dict[str, Any]]


@dataclass
class GuidelinePackage:
    guideline_text: str
    guideline_examples: list[dict[str, Any]]
    guideline_style_examples: list[dict[str, Any]]
    reference_examples: list[dict[str, Any]]
    style_catalog: list[dict[str, Any]]
    explicit_defaults: dict[str, Any]
    reference_package: DocumentPackage | None


@dataclass
class DocumentIntentProfile:
    instruction_score: int
    content_score: int
    signals: list[str]


class DocProcessor:
    def extract_document_package(self, file_bytes: bytes) -> DocumentPackage:
        document = Document(io.BytesIO(file_bytes))
        paragraphs: list[ParagraphSnapshot] = []
        tables: list[TableSnapshot] = []
        images: list[ImageSnapshot] = []
        blocks: list[dict[str, Any]] = []

        paragraph_index = 0
        table_index = 0
        image_index = 0

        for block in self._iter_block_items(document):
            if isinstance(block, Paragraph):
                image_snapshot = self._image_snapshot(
                    paragraph=block,
                    block_id=f"img-{image_index}",
                    document=document,
                )
                if image_snapshot is not None:
                    image_index += 1
                    images.append(image_snapshot)
                    blocks.append({"kind": "image", "block_id": image_snapshot.block_id})

                snapshot = self._paragraph_snapshot(
                    paragraph=block,
                    block_id=f"p-{paragraph_index}",
                )
                paragraph_index += 1
                paragraphs.append(snapshot)
                if snapshot.text:
                    blocks.append({"kind": "paragraph", "block_id": snapshot.block_id})
            elif isinstance(block, Table):
                snapshot = self._table_snapshot(
                    table=block,
                    block_id=f"t-{table_index}",
                )
                table_index += 1
                tables.append(snapshot)
                blocks.append({"kind": "table", "block_id": snapshot.block_id})

        return DocumentPackage(
            blocks=blocks,
            paragraphs=paragraphs,
            tables=tables,
            images=images,
            section=self._extract_section(document),
            style_catalog=self._extract_style_catalog(document),
        )

    def extract_guideline_package(
        self,
        guideline_bytes: bytes,
        reference_bytes: bytes | None = None,
    ) -> GuidelinePackage:
        guideline_doc = self.extract_document_package(guideline_bytes)
        reference_doc = (
            self.extract_document_package(reference_bytes) if reference_bytes else None
        )

        guideline_text = "\n".join(
            paragraph.text for paragraph in guideline_doc.paragraphs if paragraph.text
        )

        return GuidelinePackage(
            guideline_text=guideline_text,
            guideline_examples=self._build_examples(guideline_doc),
            guideline_style_examples=self._build_style_examples(guideline_doc),
            reference_examples=self._build_examples(reference_doc) if reference_doc else [],
            style_catalog=guideline_doc.style_catalog
            + (reference_doc.style_catalog if reference_doc else []),
            explicit_defaults=self._extract_explicit_defaults(guideline_text, reference_doc),
            reference_package=reference_doc,
        )

    def build_source_outline(self, source_package: DocumentPackage) -> list[dict[str, Any]]:
        outline: list[dict[str, Any]] = []
        paragraph_lookup = {item.block_id: item for item in source_package.paragraphs}
        table_lookup = {item.block_id: item for item in source_package.tables}
        image_lookup = {item.block_id: item for item in source_package.images}

        for block in source_package.blocks:
            if block["kind"] == "paragraph":
                paragraph = paragraph_lookup[block["block_id"]]
                outline.append(
                    {
                        "block_id": paragraph.block_id,
                        "kind": "paragraph",
                        "text": paragraph.text[:400],
                        "style_name": paragraph.style_name,
                        "alignment": paragraph.alignment,
                    }
                )
            else:
                if block["kind"] == "image":
                    image = image_lookup[block["block_id"]]
                    outline.append(
                        {
                            "block_id": image.block_id,
                            "kind": "image",
                            "alignment": image.alignment,
                            "width_pt": image.width_pt,
                            "height_pt": image.height_pt,
                        }
                    )
                    continue

                table = table_lookup[block["block_id"]]
                outline.append(
                    {
                        "block_id": table.block_id,
                        "kind": "table",
                        "rows": table.rows[:5],
                        "style_name": table.style_name,
                    }
                )

        return outline

    def classify_document_intent(self, document_package: DocumentPackage) -> DocumentIntentProfile:
        instruction_score = 0
        content_score = 0
        signals: list[str] = []

        combined_text = "\n".join(
            paragraph.text for paragraph in document_package.paragraphs[:120] if paragraph.text
        ).lower()

        instruction_terms = [
            "font size",
            "font emphasis",
            "alignment",
            "case",
            "line spacing",
            "paragraph spacing",
            "margins",
            "paper size",
            "table of contents",
            "heading 1",
            "heading 2",
            "heading 3",
            "cover page",
            "main body",
            "style -",
            "format of",
        ]
        content_terms = [
            "whereas",
            "now, therefore",
            "agreement",
            "party",
            "parties",
            "shall",
            "hereunder",
            "pursuant",
            "indemnity",
            "confidential",
            "termination",
            "liability",
        ]

        for term in instruction_terms:
            if term in combined_text:
                instruction_score += 2
                signals.append(f"instruction:{term}")

        for term in content_terms:
            if term in combined_text:
                content_score += 1
                signals.append(f"content:{term}")

        short_paragraph_count = 0
        long_paragraph_count = 0
        colon_count = 0
        all_caps_count = 0

        for paragraph in document_package.paragraphs[:120]:
            text = paragraph.text.strip()
            if not text:
                continue

            words = text.split()
            if len(words) <= 6:
                short_paragraph_count += 1
            if len(words) >= 25:
                long_paragraph_count += 1
            if text.endswith(":"):
                colon_count += 1
            if len(text) > 4 and text.isupper():
                all_caps_count += 1

        if short_paragraph_count >= 20:
            instruction_score += 3
            signals.append("instruction:many-short-lines")
        if colon_count >= 10:
            instruction_score += 2
            signals.append("instruction:many-colons")
        if all_caps_count >= 8:
            instruction_score += 1
            signals.append("instruction:many-all-caps-headings")
        if long_paragraph_count >= 8:
            content_score += 3
            signals.append("content:many-long-paragraphs")
        if len(document_package.tables) >= 3 and short_paragraph_count >= 15:
            instruction_score += 2
            signals.append("instruction:style-tables")

        return DocumentIntentProfile(
            instruction_score=instruction_score,
            content_score=content_score,
            signals=signals[:20],
        )

    def render_document(
        self,
        render_plan: list[dict[str, Any]],
        source_package: DocumentPackage,
        reference_package: DocumentPackage | None = None,
    ) -> DocumentType:
        document = Document()
        if reference_package and reference_package.section:
            self._apply_section_layout(document, reference_package.section)
        self._apply_document_defaults(
            document,
            render_plan,
            reference_package,
        )

        paragraph_lookup = {item.block_id: item for item in source_package.paragraphs}
        table_lookup = {item.block_id: item for item in source_package.tables}
        image_lookup = {item.block_id: item for item in source_package.images}
        plan_lookup = {item["block_id"]: item for item in render_plan}
        cover_phase = True

        for block in source_package.blocks:
            plan = plan_lookup.get(block["block_id"], {})
            if block["kind"] == "paragraph":
                paragraph_snapshot = paragraph_lookup[block["block_id"]]
                paragraph = document.add_paragraph()
                paragraph.add_run(paragraph_snapshot.text)
                self._apply_paragraph_plan(
                    paragraph,
                    paragraph_snapshot,
                    plan,
                    cover_phase=cover_phase,
                )
                if self._ends_cover_phase(paragraph_snapshot.text):
                    cover_phase = False
            elif block["kind"] == "table":
                table_snapshot = table_lookup[block["block_id"]]
                self._render_table(document, table_snapshot, plan)
            elif block["kind"] == "image":
                image_snapshot = image_lookup[block["block_id"]]
                self._render_image(document, image_snapshot)

        return document

    def _iter_block_items(self, document: DocumentType):
        parent = document.element.body
        for child in parent.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)

    def _paragraph_snapshot(self, paragraph: Paragraph, block_id: str) -> ParagraphSnapshot:
        fmt = paragraph.paragraph_format
        run_styles = [self._run_style_snapshot(run) for run in paragraph.runs if run.text]

        return ParagraphSnapshot(
            block_id=block_id,
            text=paragraph.text.strip(),
            style_name=paragraph.style.name if paragraph.style else None,
            alignment=self._alignment_to_name(paragraph.alignment),
            left_indent_pt=self._twips_to_points(fmt.left_indent),
            right_indent_pt=self._twips_to_points(fmt.right_indent),
            first_line_indent_pt=self._twips_to_points(fmt.first_line_indent),
            space_before_pt=self._twips_to_points(fmt.space_before),
            space_after_pt=self._twips_to_points(fmt.space_after),
            keep_together=fmt.keep_together,
            keep_with_next=fmt.keep_with_next,
            page_break_before=fmt.page_break_before,
            run_styles=run_styles,
        )

    def _table_snapshot(self, table: Table, block_id: str) -> TableSnapshot:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        column_count = max((len(row) for row in rows), default=0)
        return TableSnapshot(
            block_id=block_id,
            rows=rows,
            style_name=table.style.name if table.style else None,
            row_count=len(rows),
            col_count=column_count,
        )

    def _image_snapshot(
        self,
        paragraph: Paragraph,
        block_id: str,
        document: DocumentType,
    ) -> ImageSnapshot | None:
        for run in paragraph.runs:
            blips = run._r.xpath('.//*[local-name()="blip"]')
            if not blips:
                continue

            embed = blips[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not embed:
                continue

            image_part = document.part.related_parts.get(embed)
            if image_part is None:
                continue

            width_pt = None
            height_pt = None
            extents = run._r.xpath('.//*[local-name()="extent"]')
            if extents:
                cx = extents[0].get("cx")
                cy = extents[0].get("cy")
                if cx:
                    width_pt = self._emu_to_points(int(cx))
                if cy:
                    height_pt = self._emu_to_points(int(cy))

            return ImageSnapshot(
                block_id=block_id,
                image_bytes=image_part.blob,
                width_pt=width_pt,
                height_pt=height_pt,
                alignment=self._alignment_to_name(paragraph.alignment),
            )

        return None

    def _run_style_snapshot(self, run) -> RunStyleSnapshot:
        color = None
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)

        return RunStyleSnapshot(
            bold=run.bold,
            italic=run.italic,
            underline=run.underline,
            font_name=run.font.name,
            font_size_pt=run.font.size.pt if run.font.size else None,
            color_rgb=color,
            all_caps=run.font.all_caps,
        )

    def _extract_section(self, document: DocumentType) -> SectionSnapshot | None:
        if not document.sections:
            return None

        section = document.sections[0]
        return SectionSnapshot(
            top_margin_pt=self._length_to_points(section.top_margin),
            bottom_margin_pt=self._length_to_points(section.bottom_margin),
            left_margin_pt=self._length_to_points(section.left_margin),
            right_margin_pt=self._length_to_points(section.right_margin),
            page_width_pt=self._length_to_points(section.page_width),
            page_height_pt=self._length_to_points(section.page_height),
        )

    def _extract_style_catalog(self, document: DocumentType) -> list[dict[str, Any]]:
        catalog: list[dict[str, Any]] = []
        for style in document.styles:
            if not getattr(style, "name", None):
                continue
            entry = {
                "name": style.name,
                "type": str(style.type),
            }
            font = getattr(style, "font", None)
            if font:
                if font.name:
                    entry["font_name"] = font.name
                if font.size:
                    entry["font_size_pt"] = font.size.pt
                if font.bold is not None:
                    entry["bold"] = font.bold
                if font.italic is not None:
                    entry["italic"] = font.italic
            catalog.append(entry)
        return catalog[:80]

    def _build_examples(self, document_package: DocumentPackage | None) -> list[dict[str, Any]]:
        if document_package is None:
            return []

        examples: list[dict[str, Any]] = []
        for paragraph in document_package.paragraphs[:30]:
            examples.append(
                {
                    "kind": "paragraph",
                    "block_id": paragraph.block_id,
                    "text": paragraph.text[:500],
                    "style_name": paragraph.style_name,
                    "alignment": paragraph.alignment,
                    "run_styles": [asdict(run) for run in paragraph.run_styles[:3]],
                }
            )

        for table in document_package.tables[:10]:
            examples.append(
                {
                    "kind": "table",
                    "block_id": table.block_id,
                    "style_name": table.style_name,
                    "rows": table.rows[:4],
                }
            )

        return examples

    def _build_style_examples(self, document_package: DocumentPackage | None) -> list[dict[str, Any]]:
        if document_package is None:
            return []

        examples: list[dict[str, Any]] = []
        for paragraph in document_package.paragraphs[:30]:
            text = paragraph.text or ""
            examples.append(
                {
                    "kind": "paragraph",
                    "block_id": paragraph.block_id,
                    "style_name": paragraph.style_name,
                    "alignment": paragraph.alignment,
                    "word_count": len(text.split()),
                    "is_all_caps": text.isupper() if text else False,
                    "ends_with_colon": text.rstrip().endswith(":"),
                    "starts_with_numbering": self._starts_with_numbering(text),
                    "run_styles": [asdict(run) for run in paragraph.run_styles[:3]],
                }
            )

        for table in document_package.tables[:10]:
            examples.append(
                {
                    "kind": "table",
                    "block_id": table.block_id,
                    "style_name": table.style_name,
                    "row_count": table.row_count,
                    "col_count": table.col_count,
                }
            )

        return examples

    def _starts_with_numbering(self, text: str) -> bool:
        stripped = text.strip()
        if not stripped:
            return False

        prefixes = ("1.", "1)", "a.", "a)", "i.", "i)", "(a)", "(i)", "section ")
        lowered = stripped.lower()
        return lowered.startswith(prefixes)

    def _apply_section_layout(
        self,
        document: DocumentType,
        section_snapshot: SectionSnapshot,
    ) -> None:
        section = document.sections[0]
        if section_snapshot.top_margin_pt is not None:
            section.top_margin = Pt(section_snapshot.top_margin_pt)
        if section_snapshot.bottom_margin_pt is not None:
            section.bottom_margin = Pt(section_snapshot.bottom_margin_pt)
        if section_snapshot.left_margin_pt is not None:
            section.left_margin = Pt(section_snapshot.left_margin_pt)
        if section_snapshot.right_margin_pt is not None:
            section.right_margin = Pt(section_snapshot.right_margin_pt)
        if section_snapshot.page_width_pt is not None:
            section.page_width = Pt(section_snapshot.page_width_pt)
        if section_snapshot.page_height_pt is not None:
            section.page_height = Pt(section_snapshot.page_height_pt)
        section.start_type = WD_SECTION.NEW_PAGE

    def _apply_paragraph_plan(
        self,
        paragraph: Paragraph,
        source: ParagraphSnapshot,
        plan: dict[str, Any],
        cover_phase: bool = False,
    ) -> None:
        style = plan.get("style", {})
        paragraph_format = paragraph.paragraph_format

        alignment_name = style.get("alignment") or source.alignment
        if cover_phase and self._should_center_cover_line(source.text, plan):
            alignment_name = "center"
        paragraph.alignment = self._name_to_alignment(alignment_name)

        spacing_before = style.get("space_before_pt", source.space_before_pt)
        spacing_after = style.get("space_after_pt", source.space_after_pt)
        if spacing_before is not None:
            paragraph_format.space_before = Pt(spacing_before)
        if spacing_after is not None:
            paragraph_format.space_after = Pt(spacing_after)

        left_indent = style.get("left_indent_pt", source.left_indent_pt)
        right_indent = style.get("right_indent_pt", source.right_indent_pt)
        first_indent = style.get("first_line_indent_pt", source.first_line_indent_pt)
        if left_indent is not None:
            paragraph_format.left_indent = Pt(left_indent)
        if right_indent is not None:
            paragraph_format.right_indent = Pt(right_indent)
        if first_indent is not None:
            paragraph_format.first_line_indent = Pt(first_indent)

        paragraph_format.keep_together = style.get("keep_together", source.keep_together)
        paragraph_format.keep_with_next = style.get(
            "keep_with_next",
            source.keep_with_next,
        )
        paragraph_format.page_break_before = style.get(
            "page_break_before",
            source.page_break_before,
        )

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(source.text)
        if style.get("bold") is not None:
            run.bold = style["bold"]
        if style.get("italic") is not None:
            run.italic = style["italic"]
        if style.get("underline") is not None:
            run.underline = style["underline"]
        if style.get("font_name"):
            run.font.name = style["font_name"]
        if style.get("font_size_pt") is not None:
            run.font.size = Pt(style["font_size_pt"])
        if style.get("all_caps") is not None:
            run.font.all_caps = style["all_caps"]
        if style.get("color_rgb"):
            run.font.color.rgb = RGBColor.from_string(style["color_rgb"])
        elif self._is_heading_like(plan, source.text) and settings.heading_color_rgb:
            run.font.color.rgb = RGBColor.from_string(settings.heading_color_rgb)

    def _ends_cover_phase(self, text: str) -> bool:
        normalized = (text or "").strip().lower()
        if not normalized:
            return False

        if normalized == "table of contents":
            return False

        end_markers = (
            "this undertaking",
            "this agreement",
            "whereas:",
            "now therefore",
            "1. definitions",
        )
        return normalized.startswith(end_markers)

    def _should_center_cover_line(self, text: str, plan: dict[str, Any]) -> bool:
        normalized = (text or "").strip()
        lowered = normalized.lower()
        role = (plan.get("role") or "").lower()

        if not normalized:
            return False

        if role in {"title", "heading_1", "signature"}:
            return True

        center_phrases = (
            "dated ",
            "among",
            "and",
            "between",
            "table of contents",
            "as the sponsor",
            "as the borrower",
            "as the security trustee",
            "as security trustee",
        )
        if lowered in center_phrases or lowered.startswith("dated "):
            return True

        word_count = len(normalized.split())
        if word_count <= 8 and (
            normalized.isupper()
            or "limited" in lowered
            or "company limited" in lowered
            or lowered.startswith("as the ")
        ):
            return True

        return False

    def _render_table(
        self,
        document: DocumentType,
        table_snapshot: TableSnapshot,
        plan: dict[str, Any],
    ) -> None:
        if not table_snapshot.rows or table_snapshot.col_count == 0:
            return

        table = document.add_table(
            rows=len(table_snapshot.rows),
            cols=table_snapshot.col_count,
        )
        table.style = plan.get("style_name") or table_snapshot.style_name or "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        table_font = plan.get("style", {}).get("font_name") or "Times New Roman"
        table_font_size = plan.get("style", {}).get("font_size_pt") or 10

        for row_index, row in enumerate(table_snapshot.rows):
            for col_index in range(table_snapshot.col_count):
                text = row[col_index] if col_index < len(row) else ""
                cell = table.rows[row_index].cells[col_index]
                cell.text = text
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                is_header = row_index == 0 and plan.get("header_bold", True)
                if is_header and settings.table_header_fill_rgb:
                    self._shade_cell(cell, settings.table_header_fill_rgb)

                for paragraph in cell.paragraphs:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                        if is_header
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    for run in paragraph.runs:
                        run.font.name = table_font
                        run.font.size = Pt(table_font_size)
                        if is_header:
                            run.bold = True
                            if settings.table_header_font_rgb:
                                run.font.color.rgb = RGBColor.from_string(
                                    settings.table_header_font_rgb
                                )

        document.add_paragraph("")

    def _render_image(
        self,
        document: DocumentType,
        image_snapshot: ImageSnapshot,
    ) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run()

        width = Pt(image_snapshot.width_pt) if image_snapshot.width_pt else None
        height = Pt(image_snapshot.height_pt) if image_snapshot.height_pt else None

        if width and height:
            run.add_picture(io.BytesIO(image_snapshot.image_bytes), width=width, height=height)
        elif width:
            run.add_picture(io.BytesIO(image_snapshot.image_bytes), width=width)
        elif height:
            run.add_picture(io.BytesIO(image_snapshot.image_bytes), height=height)
        else:
            run.add_picture(io.BytesIO(image_snapshot.image_bytes))

        paragraph.alignment = self._name_to_alignment(
            image_snapshot.alignment if image_snapshot.alignment != "left" else "center"
        )

    def _alignment_to_name(self, alignment: WD_ALIGN_PARAGRAPH | None) -> str:
        mapping = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        return mapping.get(alignment, "left")

    def _name_to_alignment(self, alignment: str | None) -> WD_ALIGN_PARAGRAPH | None:
        mapping = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return mapping.get((alignment or "left").lower(), WD_ALIGN_PARAGRAPH.LEFT)

    def _twips_to_points(self, length) -> float | None:
        if length is None:
            return None
        return self._length_to_points(length)

    def _length_to_points(self, length) -> float | None:
        if length is None:
            return None
        return float(length.pt)

    def _emu_to_points(self, emu: int) -> float:
        return emu / 12700.0

    def _is_heading_like(self, plan: dict[str, Any], text: str) -> bool:
        role = (plan.get("role") or "").lower()
        normalized = (text or "").strip()

        if role.startswith("heading") or role in {"title", "toc", "section_heading"}:
            return True

        if not normalized:
            return False

        words = normalized.split()
        if len(words) <= 12 and (
            normalized.isupper()
            or self._starts_with_numbering(normalized)
            or normalized.lower().startswith("schedule")
        ):
            return True

        return False

    def _shade_cell(self, cell, fill_rgb: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill_rgb)

    def _extract_explicit_defaults(
        self,
        guideline_text: str,
        reference_package: DocumentPackage | None,
    ) -> dict[str, Any]:
        defaults: dict[str, Any] = {}
        lowered = guideline_text.lower()

        known_fonts = [
            "Times New Roman",
            "Arial",
            "Calibri",
            "Cambria",
            "Garamond",
            "Book Antiqua",
        ]
        for font_name in known_fonts:
            if font_name.lower() in lowered:
                defaults["font_name"] = font_name
                break

        if "single" in lowered and "line spacing" in lowered:
            defaults["line_spacing"] = "single"

        if reference_package:
            reference_font = self._extract_reference_font(reference_package)
            if reference_font and not defaults.get("font_name"):
                defaults["font_name"] = reference_font

        return defaults

    def _extract_reference_font(self, reference_package: DocumentPackage) -> str | None:
        for paragraph in reference_package.paragraphs:
            for run_style in paragraph.run_styles:
                if run_style.font_name:
                    return run_style.font_name
        return None

    def _apply_document_defaults(
        self,
        document: DocumentType,
        render_plan: list[dict[str, Any]],
        reference_package: DocumentPackage | None,
    ) -> None:
        font_name = None
        font_size_pt = None

        for item in render_plan:
            style = item.get("style", {})
            if not font_name and style.get("font_name"):
                font_name = style["font_name"]
            if not font_size_pt and style.get("font_size_pt"):
                font_size_pt = style["font_size_pt"]
            if font_name and font_size_pt:
                break

        if not font_name and reference_package:
            font_name = self._extract_reference_font(reference_package)

        normal_style = document.styles["Normal"]
        if font_name:
            normal_style.font.name = font_name
        if font_size_pt:
            normal_style.font.size = Pt(font_size_pt)
